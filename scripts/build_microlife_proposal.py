from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "MicroLife-OS_多智能体桌面生活模拟器_立项书.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "202124"
MUTED = "666666"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=None, bold=None, color=None, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        rest = text[len(bold_prefix):]
        if rest:
            r2 = p.add_run(rest)
            set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, bold=True, color=BLUE if level < 3 else DARK_BLUE, size={1: 16, 2: 13, 3: 12}[level])
    return p


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths:
            set_cell_width(cell, widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                set_cell_width(cell, widths[i])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(val)) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_font(r, size=10)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)


def add_callout(doc, label, body, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_cell_margins(table, top=120, bottom=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(label + "：")
    set_font(r, bold=True, color=DARK_BLUE)
    r2 = p.add_run(body)
    set_font(r2, color=INK)
    doc.add_paragraph()


def add_diagram(doc):
    add_table(
        doc,
        ["层级", "输入 / 处理 / 输出", "关键产物"],
        [
            ["感知层", "天气 API、系统时间、用户交互通过 IPC 写入世界事件", "外部信号事件"],
            ["决策层 MVU", "World State 维护角色、关系、记忆；Tick Scheduler 决定是否触发行为", "可回放状态变更"],
            ["LLM 适配层", "Kimi 根据人格、关系与最近上下文生成对话和动作", "结构化行为结果"],
            ["表现层", "Sprite 渲染器、多透明窗口、对话气泡组件同步展示状态", "桌面角色互动"],
            ["持久层", "SQLite 保存角色状态、关系矩阵、记忆与日志", "可恢复长期状态"],
        ],
        widths=[1800, 5000, 2560],
        header_fill=PALE_BLUE,
    )


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.167


def apply_page_setup(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def new_page(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_setup(section)
    return section


def build():
    doc = Document()
    configure_doc(doc)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MicroLife-OS")
    set_font(r, size=30, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("基于多智能体协作的桌面生活模拟沙盒")
    set_font(r, size=18, bold=True, color=DARK_BLUE)
    en = doc.add_paragraph()
    en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = en.add_run("A Multi-Agent Desktop Life Simulator")
    set_font(r, size=12, color=MUTED)
    doc.add_paragraph()
    add_callout(
        doc,
        "项目定位",
        "以大学宿舍为隐喻，构建 4 个 LLM 驱动角色共同生活、互动、记忆与演化的轻量级桌面应用原型。",
    )
    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, width_dxa=7200, indent_dxa=0)
    set_cell_margins(meta, top=100, bottom=100)
    for i, (k, v) in enumerate([
        ("项目类型", "课程实践 / 信息系统综合设计"),
        ("团队规模", "4 人"),
        ("开发周期", "14 天"),
        ("开发形态", "跨平台桌面应用（Electron）"),
        ("版本", "立项书 v1.0"),
    ]):
        meta.cell(i, 0).text = ""
        meta.cell(i, 1).text = ""
        set_cell_shading(meta.cell(i, 0), LIGHT)
        for cell, text, bold in [(meta.cell(i, 0), k, True), (meta.cell(i, 1), v, False)]:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run(text)
            set_font(rr, bold=bold)
    toc_heading = add_heading(doc, "目录", 1)
    toc_heading.paragraph_format.page_break_before = True
    add_numbers(doc, [
        "项目基本信息",
        "项目背景与意义",
        "项目目标",
        "核心功能模块",
        "技术方案",
        "创新点",
        "团队分工",
        "14 天时间规划",
        "风险评估与应对",
        "预期成果",
        "参考资料",
    ])

    first_heading = add_heading(doc, "一、项目基本信息", 1)
    first_heading.paragraph_format.page_break_before = True
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["项目名称", "MicroLife-OS：基于多智能体协作的桌面生活模拟沙盒"],
            ["英文名", "MicroLife-OS: A Multi-Agent Desktop Life Simulator"],
            ["项目类型", "课程实践 / 信息系统综合设计"],
            ["团队规模", "4 人（3 名室友 + 项目负责人）"],
            ["开发周期", "14 天"],
            ["开发形态", "跨平台桌面应用（Electron）"],
            ["关键词", "多智能体、LLM、桌面交互、状态机、人格建模"],
        ],
        widths=[2200, 7160],
    )

    add_heading(doc, "二、项目背景与意义", 1)
    add_heading(doc, "2.1 行业背景", 2)
    add_para(doc, "随着大语言模型（LLM）能力的快速增长，“AI 伙伴”类应用正从单一对话形态走向多角色交互、长期人格记忆和环境感知的新阶段。以 SillyTavern 为代表的角色卡生态，以及以 OpenAI HatchPet 为代表的桌面 AI 伙伴 Skill，共同指向一种新形态：AI 不再只是一个对话框，而是有人格、有状态、能够出现在桌面环境中的角色集合。")
    add_para(doc, "目前市场上仍缺乏一个轻量级、多智能体相互交互的桌面产品。现有方案通常偏向单宠物或单人 NPC 聊天，尚未充分支持多个 AI 角色之间形成社会关系、随时间演化，并对真实世界信号做出反应。")
    add_heading(doc, "2.2 项目立意", 2)
    add_para(doc, "本项目以大学宿舍为隐喻，将 4 个性格迥异的舍友放入桌面虚拟空间，构建一个最小但具备社会关系的生活模拟单元。项目重点验证以下问题：")
    add_numbers(doc, [
        "在有限算力与有限调用预算下，4 个 LLM 驱动角色能否产生有趣且连贯的社交互动。",
        "轻量 MVU 状态机能否替代笨重游戏引擎，承载多角色、关系与世界状态模拟。",
        "天气、时间和用户操作等真实世界数据如何自然渗入虚拟角色的行为。",
    ])
    add_heading(doc, "2.3 应用价值", 2)
    add_bullets(doc, [
        "教育层面：综合训练前后端、状态机设计、API 集成、LLM 提示工程和人机交互能力。",
        "应用层面：可延伸为桌面陪伴软件、虚拟同事、学习督促工具和心理疗愈伴侣。",
        "研究层面：为多智能体小规模社会模拟提供一个最小可运行原型。",
    ])

    add_heading(doc, "三、项目目标", 1)
    add_heading(doc, "3.1 MVP 必达目标（D14 前）", 2)
    add_table(
        doc,
        ["编号", "目标", "验收标准"],
        [
            ["M1", "4 个角色同屏共存于桌面", "同时显示、各自独立动画、可拖拽"],
            ["M2", "每个角色有独立人格", "加载 .png 角色卡，行为风格可区分"],
            ["M3", "角色之间能产生对话互动", "至少 3 种触发场景：早安、吵架、吃瓜"],
            ["M4", "接入真实天气 API", "下雨天有人抱怨，晴天有人想出门"],
            ["M5", "角色行为基于 LLM 生成", "同一情境下每次回应不同，非固定脚本"],
            ["M6", "状态持久化", "关闭再打开后，关系与记忆不丢失"],
        ],
        widths=[900, 3200, 5260],
    )
    add_heading(doc, "3.2 加分项", 2)
    add_bullets(doc, [
        "用户可以与某只角色私聊，点击角色后触发对话框。",
        "支持“今日日报”功能，晚上自动生成 4 人当天发生的事件摘要。",
        "角色之间的关系矩阵随交互自然演化，包括好感度与冲突度。",
    ])

    add_heading(doc, "四、核心功能模块", 1)
    add_diagram(doc)
    add_table(
        doc,
        ["模块", "职责", "技术"],
        [
            ["感知层", "接入外部信号", "天气 API、系统 API、IPC"],
            ["MVU 决策层", "维护世界状态与行为调度", "TypeScript 纯函数式状态机"],
            ["LLM 适配层", "调用 Kimi 生成对话与行为", "Kimi Open Platform API"],
            ["角色加载层", "解析 SillyTavern 角色卡", "PNG tEXt chunk parser"],
            ["表现层", "单窗口桌面沙盒渲染，后续可扩展透明置顶窗口", "Electron BrowserWindow + React / Canvas"],
            ["持久层", "保存状态、记忆与对话日志", "SQLite（better-sqlite3）"],
        ],
        widths=[1900, 3900, 3560],
    )

    add_heading(doc, "五、技术方案", 1)
    add_heading(doc, "5.1 技术栈", 2)
    add_table(
        doc,
        ["层级", "选型", "理由"],
        [
            ["桌面框架", "Electron 30+", "MVP 采用单 BrowserWindow 桌面沙盒，降低多窗口与 IPC 风险；后续可扩展透明置顶窗口"],
            ["前端", "React 18 + TypeScript + Vite", "组件化、类型安全、热重载快"],
            ["状态管理", "自研 MVU + Zustand", "模拟逻辑纯函数式，UI 状态轻量管理"],
            ["样式", "Tailwind CSS", "快速搭建界面，团队协作门槛低"],
            ["动画素材", "OpenAI HatchPet Skill", "一次性生成 spritesheet，运行时独立使用"],
            ["微动效", "Framer Motion", "支持气泡、淡入等过渡动效"],
            ["LLM", "Kimi（Moonshot）", "国内调用稳定，价格较低，长上下文友好"],
            ["存储", "better-sqlite3", "同步 API，本地零部署"],
            ["打包", "electron-builder", "支持一键打包 dmg / exe"],
        ],
        widths=[1600, 3000, 4760],
    )

    add_heading(doc, "5.2 关键技术决策", 2)
    add_heading(doc, "① 桌面架构策略", 3)
    add_para(doc, "项目仍采用 Electron 作为桌面容器，但 MVP 阶段优先实现单 BrowserWindow 内的多角色桌面沙盒：4 个角色以可拖拽组件形式存在于同一渲染窗口中，由统一状态机驱动位置、动作和对话气泡。这样既保留“桌面应用”的项目定位，又能显著降低多窗口同步、窗口层级、跨平台权限和打包调试风险。")
    add_para(doc, "当核心闭环稳定后，再将透明置顶窗口或多 BrowserWindow 桌宠形态作为加分扩展，而不是 MVP 的强依赖。")
    add_heading(doc, "② 美术资产策略", 3)
    add_para(doc, "创作阶段使用 HatchPet Skill 生成 spritesheet.webp 与 pet.json；运行阶段由自研 Sprite 渲染器读取 assets/chars/*.webp 与 assets/chars/*.json，不再依赖 Codex 或 HatchPet 环境。")
    add_heading(doc, "③ 多智能体协作模式", 3)
    add_bullets(doc, [
        "采用 Tick 驱动 + 事件总线模式，全局 tick 频率为 6 秒/帧，平衡 LLM 成本与互动密度。",
        "每个 tick 评估角色是否发起行动，依据人格、当前状态和关系矩阵决定行为。",
        "行动触发 LLM 调用，生成结果回写世界状态；极端情境通过事件总线即时触发。",
    ])
    add_heading(doc, "④ 角色卡兼容性", 3)
    add_bullets(doc, [
        "直接读取 SillyTavern v2 PNG 角色卡，解析 tEXt chunk 中嵌入的 JSON。",
        "用户可从角色卡社区下载现成角色，降低造人成本，项目聚焦于造社会。",
    ])
    add_heading(doc, "5.3 LLM 调用预算控制", 2)
    add_table(
        doc,
        ["措施", "效果"],
        [
            ["6 秒 tick + 概率触发", "平均每分钟约 5-10 次调用"],
            ["短上下文", "仅传当前情境与最近 3 轮，单次调用 < 1500 tokens"],
            ["关键事件全 LLM，日常动作走模板", "成本预计降低约 60%"],
            ["两周开发期预估成本", "< 30 元 RMB"],
        ],
        widths=[4300, 5060],
    )

    add_heading(doc, "六、创新点", 1)
    add_table(
        doc,
        ["编号", "创新点", "区别于现有方案"],
        [
            ["1", "多智能体桌面社交模拟", "不同于单宠物或单聊，本项目让 4 个 LLM 角色互相对话"],
            ["2", "MVU 状态机驱动 LLM 行为", "状态机决定何时调用 LLM，成本可控、行为可预测"],
            ["3", "现实世界感知", "天气、时间与用户行为渗入虚拟角色，形成虚实弱耦合"],
            ["4", "生态复用", "复用 SillyTavern 角色卡与 HatchPet sprite 生成能力"],
            ["5", "轻量、可演示、可商业化", "本地运行，低云依赖，安装包体积可控"],
        ],
        widths=[900, 3000, 5460],
    )

    add_heading(doc, "七、团队分工", 1)
    add_para(doc, "本团队由 4 名大一学生组成。项目负责人承担大部分框架搭建、核心逻辑和集成兜底；其余 3 名成员技术水平接近，主要负责可拆分、可验证、风险较低的模块，并在联调阶段共同参与测试与素材补全。")
    add_table(
        doc,
        ["角色", "负责模块", "主要任务"],
        [
            ["项目负责人", "项目框架 / MVU / Kimi 集成 / 总联调 / 答辩", "搭建 Electron + React + TypeScript 项目骨架；设计 World State、Tick Scheduler 与事件总线；接入 Kimi API、SQLite 与天气数据；负责最终集成、演示脚本和答辩主讲"],
            ["成员 A", "角色素材 / 人设配置 / Sprite 资源整理", "整理 4 个角色的人设、口头禅和性格标签；使用或整理 HatchPet 产物；维护 assets/chars 资源目录；配合检查角色动画与展示效果"],
            ["成员 B", "界面组件 / 对话气泡 / 基础交互", "在项目负责人给出的框架内完成角色卡片、对话气泡、按钮、设置项等 UI 组件；实现基础点击、拖拽和私聊弹窗的页面部分"],
            ["成员 C", "天气配置 / 测试记录 / 文档材料", "整理天气 API 字段与配置说明；维护测试用例和开发日志；协助 README、部署文档、结题报告素材和演示视频脚本"],
        ],
        widths=[1500, 3200, 4660],
    )
    add_para(doc, "分工原则为“框架集中、任务拆小、每日可验收”。项目负责人保证主线能跑通，其他成员围绕素材、界面、配置、测试和文档持续交付，关键节点（D7 联调、D12 集成测试）全员参与。")

    add_heading(doc, "八、14 天时间规划", 1)
    add_table(
        doc,
        ["阶段", "天数", "关键产出"],
        [
            ["筹备期", "D1-D2", "立项书定稿、Electron 骨架、4 套 HatchPet 资产"],
            ["核心期", "D3-D7", "MVU 世界模型、Kimi 接通、角色卡加载、单角色可桌面活动"],
            ["协作期", "D8-D10", "单窗口 4 角色同屏、天气接入、多角色对话、状态持久化"],
            ["打磨期", "D11-D13", "性能优化、集成测试、打包、演示脚本"],
            ["答辩", "D14", "演示、答辩、报告提交"],
        ],
        widths=[1700, 1400, 6260],
    )
    add_table(
        doc,
        ["D1", "D2", "D3", "D4", "D5", "D6", "D7", "D8", "D9", "D10", "D11", "D12", "D13", "D14"],
        [["立项", "资产", "骨架", "MVU", "LLM", "角色卡", "联调 M1", "4 角色同屏", "天气 / 日报", "持久化", "优化", "集成测试", "打包", "答辩"]],
        widths=[668] * 14,
        header_fill=PALE_BLUE,
    )
    add_callout(doc, "关键里程碑", "D7 完成单角色闭环；D10 完成单窗口 4 角色同屏与至少 1 个完整互动场景；D13 形成可打包分发安装包。透明置顶或多窗口形态作为稳定后的扩展目标。")

    add_heading(doc, "九、风险评估与应对", 1)
    add_table(
        doc,
        ["编号", "风险", "概率", "影响", "应对策略"],
        [
            ["R1", "HatchPet 产物格式不兼容，无法独立渲染", "中", "高", "D2 立即验证，失败则切换 Rive 兜底方案"],
            ["R2", "Kimi API 限流或偶发失败", "中", "中", "重试、模板回退、本地缓存对话池"],
            ["R3", "Electron 多窗口与透明窗口调试复杂", "中", "中", "MVP 采用单 BrowserWindow 多角色组件方案，多窗口仅作为加分扩展"],
            ["R4", "4 角色对话失控或不自然", "高", "中", "系统提示词约束、关系矩阵硬限制、人工调教"],
            ["R5", "团队成员进度跟不上", "中", "高", "每日早晚站会，任务卡拆分，关键模块负责人冗余"],
            ["R6", "LLM 成本超预算", "低", "低", "tick 频率可降，模板回退，设置硬上限"],
            ["R7", "跨平台打包问题", "中", "低", "主推 macOS 或 Windows 任一平台，另一个标注为实验性"],
        ],
        widths=[700, 3300, 800, 800, 3760],
    )

    add_heading(doc, "十、预期成果", 1)
    add_heading(doc, "10.1 交付物", 2)
    add_numbers(doc, [
        "可运行的 macOS / Windows 安装包。",
        "完整源代码、README 与部署文档。",
        "项目结题报告（技术、设计、用户测试）。",
        "5 分钟演示视频。",
        "答辩 PPT。",
    ])
    add_heading(doc, "10.2 答辩亮点", 2)
    add_callout(doc, "亮点一", "本项目站在 OpenAI HatchPet 与 SillyTavern 两个生态之上，将其尚未解决的多智能体桌面社交模拟问题，在 14 天内做成可运行原型。", fill=LIGHT)
    add_callout(doc, "亮点二", "项目用 MVU 状态机把 LLM 从黑盒生成器转化为可控行为引擎，使 4 个角色能在每分钟不到 10 次调用的预算下产生连贯社交行为。", fill=LIGHT)
    add_callout(doc, "亮点三", "项目不仅是课程作业，也指向 LLM 时代的新产品形态：有人格、有关系、住在桌面上的 AI 角色集合。", fill=LIGHT)

    add_heading(doc, "十一、参考资料", 1)
    add_numbers(doc, [
        "OpenAI Skills Repository — HatchPet Skill 源码与文档。",
        "SillyTavern Character Card V2 Specification。",
        "Moonshot Kimi Open Platform API Docs。",
        "Electron 官方文档 — BrowserWindow 与 IPC。",
        "The Elm Architecture（MVU 模式起源）。",
    ])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = footer.add_run("MicroLife-OS 立项书")
    set_font(rr, size=9, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
